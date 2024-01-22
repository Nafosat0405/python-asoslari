from docx import Document
class Word:
    def __init__(self):
        self.hujjat=Document()
        
    def matn_qush(self,matn):
        self.hujjat.add_paragraph(matn)
        
    def sarlavha_qush(self,sarlavha):
        self.hujjat.add_heading(sarlavha)
        
    def saqlash(self):
         self.hujjat.save("hujjatlar/matnlar_tuplami.docx")
    

obj1=Word()
obj1.matn_qush("salom")
obj1.sarlavha_qush("Uzbekistan")
obj1.saqlash()